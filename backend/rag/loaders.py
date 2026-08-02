"""Document loaders for PDF, text, Word, and CSV."""

from pathlib import Path
from typing import List

from langchain_core.documents import Document


def load_text(path: str | Path) -> List[Document]:
    p = Path(path)
    content = p.read_text(encoding="utf-8", errors="ignore")
    return [Document(page_content=content, metadata={"source": str(p), "type": "text"})]


def load_pdf(path: str | Path) -> List[Document]:
    from pypdf import PdfReader

    p = Path(path)
    reader = PdfReader(str(p))
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        docs.append(
            Document(
                page_content=text,
                metadata={"source": str(p), "type": "pdf", "page": i + 1},
            )
        )
    return docs


def load_docx(path: str | Path) -> List[Document]:
    from docx import Document as DocxDocument

    p = Path(path)
    doc = DocxDocument(str(p))
    text = "\n".join(para.text for para in doc.paragraphs)
    return [Document(page_content=text, metadata={"source": str(p), "type": "docx"})]


def load_csv(path: str | Path) -> List[Document]:
    import pandas as pd

    p = Path(path)
    df = pd.read_csv(p)
    docs: List[Document] = []
    for idx, row in df.iterrows():
        content = "\n".join(f"{k}: {v}" for k, v in row.to_dict().items())
        docs.append(
            Document(
                page_content=content,
                metadata={"source": str(p), "type": "csv", "row": int(idx)},
            )
        )
    return docs


def load_document(path: str | Path) -> List[Document]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(p)
    if suffix in {".docx", ".doc"}:
        return load_docx(p)
    if suffix == ".csv":
        return load_csv(p)
    return load_text(p)
